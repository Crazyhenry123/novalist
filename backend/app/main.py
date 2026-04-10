"""FastAPI application — backend for Novalist v2 step-based novel composition."""

import logging
from contextlib import asynccontextmanager

from fastapi import FastAPI, Request, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse

from app.config import settings
from app.models.schemas import (
    Step1Request,
    Step2Request,
    Step3ChapterRequest,
    ChatRequest,
    SaveStep1Request,
    SaveStep2Request,
)
from app.models.novel_store import NovelStore
from app.agents.orchestrator import (
    stream_step1,
    stream_step2,
    stream_step3_chapter,
    stream_chat,
)

logging.basicConfig(level=getattr(logging, settings.log_level))
logger = logging.getLogger(__name__)


@asynccontextmanager
async def lifespan(app: FastAPI):
    logger.info("Novalist backend starting")
    yield
    logger.info("Novalist backend shutting down")


app = FastAPI(title="Novalist", version="2.0.0", lifespan=lifespan)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=["Content-Type"],
)


# ── Health ───────────────────────────────────────────────────────────

@app.get("/health")
async def health():
    return {"status": "ok", "service": "novalist"}


# ── SSE streaming endpoints ─────────────────────────────────────────

def _get_user_id(request: Request) -> str:
    return request.query_params.get("user_id", "anonymous")


@app.post("/api/step1")
async def step1(request: Request):
    """Stream step 1: generate structure, characters, world in parallel."""
    body = await request.json()
    user_id = _get_user_id(request)
    logger.info("Step 1 request from user %s", user_id)

    try:
        req = Step1Request(**body)
    except Exception as e:
        logger.error("Invalid step1 request: %s", e)
        return {"status": "error", "message": f"请求无效：{e}"}

    return StreamingResponse(
        stream_step1(req, user_id),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache, no-transform",
            "X-Accel-Buffering": "no",
        },
    )


@app.post("/api/step2")
async def step2(request: Request):
    """Stream step 2: generate plot outline from step 1 results."""
    body = await request.json()
    user_id = _get_user_id(request)
    logger.info("Step 2 request from user %s", user_id)

    try:
        req = Step2Request(**body)
    except Exception as e:
        logger.error("Invalid step2 request: %s", e)
        return {"status": "error", "message": f"请求无效：{e}"}

    return StreamingResponse(
        stream_step2(req, user_id),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache, no-transform",
            "X-Accel-Buffering": "no",
        },
    )


@app.post("/api/step3/chapter")
async def step3_chapter(request: Request):
    """Stream step 3: write and edit one chapter."""
    body = await request.json()
    user_id = _get_user_id(request)
    logger.info("Step 3 chapter request from user %s", user_id)

    try:
        req = Step3ChapterRequest(**body)
    except Exception as e:
        logger.error("Invalid step3 request: %s", e)
        return {"status": "error", "message": f"请求无效：{e}"}

    return StreamingResponse(
        stream_step3_chapter(req, user_id),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache, no-transform",
            "X-Accel-Buffering": "no",
        },
    )


@app.post("/api/chat")
async def chat(request: Request):
    """Stream a brainstorming chat response."""
    body = await request.json()
    user_id = _get_user_id(request)
    logger.info("Chat request from user %s", user_id)

    try:
        req = ChatRequest(**body)
    except Exception as e:
        logger.error("Invalid chat request: %s", e)
        return {"status": "error", "message": f"请求无效：{e}"}

    return StreamingResponse(
        stream_chat(req, user_id),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache, no-transform",
            "X-Accel-Buffering": "no",
        },
    )


# ── REST endpoints for saving edits and reading state ────────────────

@app.put("/api/novel/{novel_id}/step1")
async def save_step1(novel_id: str, request: Request):
    """Save user edits to step 1 results."""
    body = await request.json()
    user_id = _get_user_id(request)
    logger.info("Save step1 for novel %s, user %s", novel_id, user_id)

    try:
        req = SaveStep1Request(**body)
    except Exception as e:
        return {"status": "error", "message": f"请求无效：{e}"}

    store = NovelStore()
    store.save_step1(user_id, novel_id, req.structure, req.characters, req.world)
    novel = store.get_novel_full(user_id, novel_id)
    return novel


@app.put("/api/novel/{novel_id}/step2")
async def save_step2(novel_id: str, request: Request):
    """Save user edits to step 2 results."""
    body = await request.json()
    user_id = _get_user_id(request)
    logger.info("Save step2 for novel %s, user %s", novel_id, user_id)

    try:
        req = SaveStep2Request(**body)
    except Exception as e:
        return {"status": "error", "message": f"请求无效：{e}"}

    store = NovelStore()
    store.save_step2(user_id, novel_id, req.plot)
    novel = store.get_novel_full(user_id, novel_id)
    return novel


@app.get("/api/novels")
async def list_novels(request: Request):
    """List all novels for a user (DynamoDB index only, no S3)."""
    user_id = _get_user_id(request)
    store = NovelStore()
    novels = store.list_novels(user_id)
    return {"novels": novels}


@app.get("/api/novel/{novel_id}")
async def get_novel(novel_id: str, request: Request):
    """Get full novel state including S3 content and chapters."""
    user_id = _get_user_id(request)
    store = NovelStore()
    novel = store.get_novel_full(user_id, novel_id)
    if not novel:
        return {"status": "error", "message": "小说不存在"}
    return novel


@app.delete("/api/novel/{novel_id}")
async def delete_novel(novel_id: str, request: Request):
    """Delete a novel and all its S3 content."""
    user_id = _get_user_id(request)
    logger.info("Delete novel %s for user %s", novel_id, user_id)
    store = NovelStore()
    store.delete_novel(user_id, novel_id)
    return {"status": "ok", "message": "已删除"}


# ── Memory endpoints ─────────────────────────────────────────────────

@app.get("/api/memory")
async def get_memory(request: Request, novel_id: str = Query(..., description="小说ID")):
    """Get the novel's session memory."""
    user_id = _get_user_id(request)
    from app.models.memory import MemoryManager
    mgr = MemoryManager(user_id, novel_id)
    return mgr.load()


@app.put("/api/memory")
async def update_memory(request: Request):
    """Update the novel's session memory."""
    user_id = _get_user_id(request)
    body = await request.json()
    novel_id = body.get("novel_id")
    if not novel_id:
        return {"status": "error", "message": "缺少 novel_id"}
    from app.models.memory import MemoryManager
    mgr = MemoryManager(user_id, novel_id)
    memory = mgr.load()
    if "user_preferences" in body:
        memory["user_preferences"] = body["user_preferences"]
    if "current_novel" in body:
        memory["current_novel"].update(body["current_novel"])
    mgr.save(memory)
    return {"status": "ok"}


# ── Mark complete ──────────────────────────────────────────────────

@app.post("/api/novel/{novel_id}/complete")
async def mark_complete(novel_id: str, request: Request):
    """Mark a novel as completed."""
    user_id = _get_user_id(request)
    store = NovelStore()
    store.update_status(user_id, novel_id, "completed")
    return {"status": "ok", "message": "已标记为完成"}


# ── Export endpoints (docx) ────────────────────────────────────────

def _build_docx(title: str, sections: list[tuple[str, str]]) -> bytes:
    """Build a docx document from sections [(heading, body), ...]."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io

    doc = Document()

    # Title
    t = doc.add_heading(title, level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for heading, body in sections:
        if heading:
            doc.add_heading(heading, level=1)
        if body:
            for para_text in body.split("\n"):
                stripped = para_text.strip()
                if not stripped:
                    continue
                # Sub-headings: lines starting with 【】
                if stripped.startswith("【") and "】" in stripped:
                    h = doc.add_heading(stripped, level=2)
                else:
                    p = doc.add_paragraph(stripped)
                    p.style.font.size = Pt(11)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@app.get("/api/novel/{novel_id}/export")
async def export_novel(novel_id: str, request: Request):
    """Export full novel as a docx document."""
    user_id = _get_user_id(request)
    store = NovelStore()
    novel = store.get_novel_full(user_id, novel_id)
    if not novel:
        return {"status": "error", "message": "小说不存在"}

    title = novel.get("premise", "未命名小说")[:100]
    sections = []

    if novel.get("structure"):
        sections.append(("故事结构", novel["structure"]))
    if novel.get("characters"):
        sections.append(("角色设定", novel["characters"]))
    if novel.get("world"):
        sections.append(("世界观", novel["world"]))
    if novel.get("plot"):
        sections.append(("情节大纲", novel["plot"]))

    chapters = novel.get("chapters", {})
    if chapters:
        for num in sorted(chapters.keys(), key=lambda x: int(x)):
            sections.append((f"第{num}章", chapters[num]))

    docx_bytes = _build_docx(title, sections)

    from fastapi.responses import Response
    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f'attachment; filename="novel-{novel_id[:8]}.docx"',
        },
    )


@app.get("/api/novel/{novel_id}/chapter/{chapter_num}/export")
async def export_chapter(novel_id: str, chapter_num: int, request: Request):
    """Export a single chapter as docx."""
    user_id = _get_user_id(request)
    store = NovelStore()
    content = store.s3.load_text(f"users/{user_id}/novels/{novel_id}/chapters/{chapter_num:02d}.md")
    if not content:
        return {"status": "error", "message": "章节不存在"}

    docx_bytes = _build_docx(f"第{chapter_num}章", [("", content)])

    from fastapi.responses import Response
    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f'attachment; filename="chapter-{chapter_num:02d}.docx"',
        },
    )
